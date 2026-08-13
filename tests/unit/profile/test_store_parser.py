"""档案 YAML 存取与简历解析单元测试（离线，LLM 用 fake）。"""

from __future__ import annotations

from pathlib import Path

import pytest

from autooffer_core.errors import ProfileError
from autooffer_core.profile.parser import extract_text, parse_resume
from autooffer_core.profile.schema import Profile
from autooffer_core.profile.store import load_profile, profile_from_yaml, profile_to_yaml
from autooffer_core.testing import FakeLLMClient, build_sample_profile


def test_yaml_roundtrip() -> None:
    p = build_sample_profile()
    text = profile_to_yaml(p)
    restored = profile_from_yaml(text)
    assert restored == p


def test_yaml_invalid_raises() -> None:
    with pytest.raises(ProfileError):
        profile_from_yaml("basic: [不是合法档案")


def test_load_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(ProfileError):
        load_profile(tmp_path / "nope.yaml")


def test_extract_text_txt(tmp_path: Path) -> None:
    f = tmp_path / "resume.txt"
    f.write_text("张三 电话 13800001111", encoding="utf-8")
    assert "13800001111" in extract_text(str(f))


def test_extract_text_docx(tmp_path: Path) -> None:
    import docx

    doc = docx.Document()
    doc.add_paragraph("李四的简历")
    doc.add_paragraph("邮箱: lisi@example.com")
    f = tmp_path / "resume.docx"
    doc.save(str(f))
    text = extract_text(str(f))
    assert "李四的简历" in text
    assert "lisi@example.com" in text


def test_extract_text_unsupported(tmp_path: Path) -> None:
    f = tmp_path / "resume.png"
    f.write_bytes(b"\x89PNG")
    with pytest.raises(ProfileError):
        extract_text(str(f))


@pytest.mark.asyncio
async def test_parse_resume_with_fake_llm(tmp_path: Path) -> None:
    f = tmp_path / "resume.txt"
    f.write_text("张三，本科，示例大学计算机专业", encoding="utf-8")

    fake_profile = build_sample_profile().model_copy(update={"id": "auto"})
    payload = (
        '{"profile": ' + fake_profile.model_dump_json() + ', "low_confidence_paths": '
        '["basic.birth_date"]}'
    )
    llm = FakeLLMClient(payload)

    profile, low_conf = await parse_resume(str(f), llm)
    assert isinstance(profile, Profile)
    assert profile.id.startswith("profile-")  # auto 被替换为真实 id
    assert profile.basic.name == "张三"
    assert low_conf == ["basic.birth_date"]
    # 来源简历自动登记为附件，供"上传简历"字段直接使用
    resumes = [a for a in profile.attachments if a.kind == "resume"]
    assert resumes
    assert any(str(f) in a.path or a.path.endswith(f.name) for a in resumes)
    # LLM 收到的提示词包含简历全文与 schema
    sent = llm.messages_seen[0][0].content
    assert "示例大学计算机专业" in sent
    assert "low_confidence_paths" in sent
