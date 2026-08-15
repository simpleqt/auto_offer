"""简历文件解析为结构化档案（docs/03 §1.2，FR-P1/P9）。

PDF 用 pypdf、Word 用 python-docx 抽文本，LLM 按 Profile schema 结构化抽取；
返回 (档案, 低置信字段路径列表)，低置信字段由界面高亮请用户确认。
"""

from __future__ import annotations

import asyncio
import json
import uuid
from pathlib import Path

import structlog
from pydantic import BaseModel, Field

from autooffer_core.errors import ProfileError
from autooffer_core.llm.interfaces import ChatMessage, LLMClient
from autooffer_core.profile.schema import Attachment, Profile

log = structlog.get_logger(__name__)

_MAX_TEXT_CHARS = 30_000


class _ExtractedProfile(BaseModel):
    """LLM 抽取输出：档案主体 + 低置信字段路径。"""

    profile: Profile
    low_confidence_paths: list[str] = Field(default_factory=list)


def _extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(file_path: str) -> str:
    """从简历文件抽取纯文本（同步；上层用 to_thread 包装）。"""
    p = Path(file_path)
    if not p.is_file():
        raise ProfileError(f"简历文件不存在: {file_path}")
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf_text(p)
    elif suffix in (".docx", ".doc"):
        if suffix == ".doc":
            raise ProfileError("暂不支持旧版 .doc，请另存为 .docx 或 PDF")
        text = _extract_docx_text(p)
    elif suffix in (".txt", ".md"):
        text = p.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ProfileError(f"不支持的简历格式: {suffix}（支持 pdf/docx/txt/md）")
    text = text.strip()
    if not text:
        raise ProfileError("简历文本抽取为空（可能是扫描件，请提供文字版）")
    return text[:_MAX_TEXT_CHARS]


_PARSE_PROMPT = """你是简历信息抽取器。请从下面的简历全文中抽取结构化档案，严格按 JSON Schema 输出。

要求：
1. 只抽取简历中明确存在的信息，绝不编造；没有的字段留空/省略。
2. 日期解析为 {{"year": 2023, "month": 7, "day": 15}} 形式：原文写到"日"就必须填 day
   （如"2001年3月18日出生" → {{"year": 2001, "month": 3, "day": 18}}），只写到月则省略 day；
   "至今"的经历 end 置 null。
3. 经历 kind：实习填 internship、正式工作填 work、项目填 project。
4. 获奖、语言等级（四六级/雅思等）、社团/学生干部经历放入 extended 对应字段。
5. profile.id 填 "auto"，profile.label 填 "解析-待确认"。
6. 对拿不准的字段（字迹模糊、表述含糊、你做了推断的），把字段路径放入 low_confidence_paths，
   如 ["basic.birth_date", "education[0].gpa"]。

JSON Schema：
{schema}

简历全文：
---
{resume_text}
---

只输出一个 JSON 对象：{{"profile": {{...}}, "low_confidence_paths": [...]}}"""


def _resume_attachment(file_path: str, text: str) -> Attachment:
    """把来源简历文件登记为档案附件，供表单的"上传简历"字段直接使用。"""
    is_zh = any("\u4e00" <= ch <= "\u9fff" for ch in text[:2000])
    return Attachment(
        kind="resume",
        label="中文简历" if is_zh else "英文简历",
        path=str(Path(file_path).resolve()),
        language="zh" if is_zh else "en",
    )


async def parse_resume(file_path: str, llm: LLMClient) -> tuple[Profile, list[str]]:
    """解析简历文件 → (Profile, 低置信字段路径)。

    来源简历文件会自动登记为档案附件（kind=resume），
    这样"上传简历"类字段无需用户再手工配置附件路径。
    """
    text = await asyncio.to_thread(extract_text, file_path)
    schema_json = json.dumps(Profile.model_json_schema(), ensure_ascii=False)
    prompt = _PARSE_PROMPT.format(schema=schema_json, resume_text=text)
    result = await llm.complete_json(
        [ChatMessage(role="user", content=prompt)], _ExtractedProfile
    )
    assert isinstance(result, _ExtractedProfile)
    profile = result.profile
    if not profile.id or profile.id == "auto":
        profile = profile.model_copy(update={"id": f"profile-{uuid.uuid4().hex[:8]}"})
    # 来源文件是确定事实：按路径判重后登记（不因模型臆造的附件项而跳过）
    source = _resume_attachment(file_path, text)
    if not any(Path(a.path) == Path(source.path) for a in profile.attachments):
        profile.attachments.append(source)
    log.info(
        "resume.parsed",
        file=Path(file_path).name,
        educations=len(profile.education),
        experiences=len(profile.experiences),
        low_confidence=len(result.low_confidence_paths),
    )
    return profile, result.low_confidence_paths
